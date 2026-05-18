#!/usr/bin/env python3
"""
Build docs/ServeFlow-Documentation.docx from docs/ServeFlow-Documentation.md.

Usage (from repo root):
    pip install python-docx
    python docs/generate_docx.py
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("python-docx is required: pip install python-docx", file=sys.stderr)
    sys.exit(1)

DOCS_DIR = Path(__file__).resolve().parent
MD_PATH = DOCS_DIR / "ServeFlow-Documentation.md"
OUT_PATH = DOCS_DIR / "ServeFlow-Documentation.docx"

CODE_FENCE = re.compile(r"^```(\w*)")
HEADING = re.compile(r"^(#{1,6})\s+(.*)")
TABLE_ROW = re.compile(r"^\|(.+)\|$")
TABLE_SEP = re.compile(r"^\|[\s\-:|]+\|$")
BOLD_INLINE = re.compile(r"\*\*(.+?)\*\*")
MERMAID_START = re.compile(r"^```mermaid")


def add_formatted_paragraph(doc: Document, text: str, style: str | None = None):
    p = doc.add_paragraph(style=style)
    pos = 0
    for m in BOLD_INLINE.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos : m.start()])
        run = p.add_run(m.group(1))
        run.bold = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])
    return p


def parse_table_rows(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines):
        line = lines[i].strip()
        if not TABLE_ROW.match(line):
            break
        if TABLE_SEP.match(line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def md_to_docx(md_text: str) -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_lang = ""
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not in_code and MERMAID_START.match(stripped):
            in_code = True
            code_lang = "mermaid"
            code_lines = []
            i += 1
            continue

        if not in_code:
            m_fence = CODE_FENCE.match(stripped)
            if m_fence:
                in_code = True
                code_lang = m_fence.group(1) or "text"
                code_lines = []
                i += 1
                continue

        if in_code:
            if stripped == "```":
                label = "Diagram (Mermaid — paste into Word or use mermaid.live):" if code_lang == "mermaid" else f"Code ({code_lang}):"
                p = doc.add_paragraph()
                run = p.add_run(label)
                run.bold = True
                for cl in code_lines:
                    cp = doc.add_paragraph(cl)
                    cp.style = "Intense Quote"
                in_code = False
                code_lines = []
                i += 1
                continue
            code_lines.append(line)
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph("—" * 40)
            i += 1
            continue

        hm = HEADING.match(line)
        if hm:
            level = len(hm.group(1))
            title = hm.group(2).strip()
            if level == 1:
                doc.add_heading(title, level=0)
            else:
                doc.add_heading(title, level=min(level, 3))
            i += 1
            continue

        if TABLE_ROW.match(stripped):
            rows, i = parse_table_rows(lines, i)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        if ci < len(table.rows[ri].cells):
                            table.rows[ri].cells[ci].text = cell
                doc.add_paragraph()
            continue

        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", stripped), style="List Number")
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        add_formatted_paragraph(doc, stripped)
        i += 1

    return doc


def main() -> int:
    if not MD_PATH.is_file():
        print(f"Missing source file: {MD_PATH}", file=sys.stderr)
        return 1

    md_text = MD_PATH.read_text(encoding="utf-8")
    document = md_to_docx(md_text)
    document.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
